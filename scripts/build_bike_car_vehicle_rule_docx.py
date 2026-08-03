from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "bike-car-driver-vehicle-rule-flow.docx"
ASSETS = ROOT / "docs" / "assets"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(85, 85, 85)
LIGHT_GRAY = "F2F4F7"
BORDER = "D9E2F3"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="DADCE0", sz="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), sz)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx >= len(row.cells):
                continue
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def style_table(table, widths):
    set_table_widths(table, widths)
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_GRAY)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = DARK_BLUE


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(subtitle)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.style = "Heading 1"
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = BLUE


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.style = "Heading 2"
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = BLUE


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.12
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.name = "Calibri"
        r2 = p.add_run(text[len(bold_prefix):])
        r2.font.size = Pt(10.5)
        r2.font.name = "Calibri"
    else:
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        r.font.name = "Calibri"
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.1
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(10)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.1
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(10)


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    style_table(table, [6.3])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F8FF")
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    r.font.size = Pt(10.5)
    r2 = p.add_run(f"\n{text}")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(40, 40, 40)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        table.cell(0, i).text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    style_table(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_image(doc, filename, caption):
    path = ASSETS / filename
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(caption)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    r.font.size = Pt(10.5)
    doc.add_picture(str(path), width=Inches(6.35))
    pic_p = doc.paragraphs[-1]
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.space_after = Pt(8)


def add_step_flow(doc, title, steps):
    add_h2(doc, title)
    for idx, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{idx}. ")
        r.bold = True
        r.font.color.rgb = DARK_BLUE
        r.font.size = Pt(10.5)
        r2 = p.add_run(step)
        r2.font.size = Pt(10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    add_title(
        doc,
        "Bike/Car — Quy trình loại xe, giá tiền và gán tài xế",
        "Tài liệu Word dùng cho BA/CMS/BE/FE đọc luồng vận hành, có nhúng ảnh minh họa để mở trực tiếp trên Word.",
    )

    add_callout(
        doc,
        "Nguyên tắc chốt",
        "Bike/Car không tách tài xế và phương tiện thành hai nhóm riêng. Với Bike/Car: 1 tài xế = 1 xe riêng. Hồ sơ tài xế chỉ cần chọn nhóm xe, tick nhanh Ô tô 6–7 chỗ/Premium nếu có; hệ thống tự suy ra loại xe được phép nhận.",
    )

    add_h1(doc, "1. Màn hình minh họa")
    add_image(doc, "bike-car-vehicle-type-list.png", "Hình 1 — Danh sách Loại xe trong Dữ liệu vận tải")
    add_body(doc, "Màn hình Loại xe dùng để tạo các loại xe khách có thể chọn khi đặt Bike/Car: Bike phổ thông, Bike Premium, Car 04 phổ thông, Car 4 Premium, Car 06 phổ thông, Car 06 Premium.")
    doc.add_page_break()
    add_image(doc, "bike-car-pricing-tab.png", "Hình 2 — Giá tiền theo tab Xe máy/Xe hơi")
    add_body(doc, "Giá tiền được lọc theo tab Xe máy hoặc Xe hơi. Admin tạo bảng giá theo từng loại xe, không tạo giá chung mơ hồ.")
    add_image(doc, "bike-car-pricing-modal.png", "Hình 3 — Modal tạo bảng giá cho loại xe")
    add_body(doc, "Khi tạo bảng giá, danh sách loại xe áp dụng được lọc theo tab đang mở. Tab Xe máy chỉ chọn loại xe máy; tab Xe hơi chỉ chọn loại xe hơi.")
    doc.add_page_break()

    add_h1(doc, "2. Phân vai màn hình")
    add_table(
        doc,
        ["Nhóm", "Màn hình", "Mục đích"],
        [
            ["CMS", "Loại xe", "Tạo/sửa/tạm dừng loại xe Bike/Car để khách chọn"],
            ["CMS", "Giá tiền", "Tạo giá theo từng loại xe"],
            ["CMS", "Nhà xe & Tài xế", "Chọn nhóm xe, tick nhanh Ô tô 6–7 chỗ/Premium và xem loại xe hệ thống cho phép nhận"],
            ["CMS", "Nhiệm vụ phân công", "Lọc tài xế available và điều phối thủ công khi cần"],
            ["BE", "Master data + Matching", "Lưu loại xe, giá, rule quyền nhận và chỉ matching đúng tài xế"],
            ["FE Customer", "Booking Bike/Car", "Hiển thị loại xe hoạt động, có giá và có tài xế available"],
            ["FE Driver", "App tài xế", "Online/Offline, nhận offer đúng loại xe được phép"],
            ["BA", "Nghiệp vụ", "Chốt rule Ô tô 6–7 chỗ, Premium, trạng thái và UAT"],
        ],
        [0.8, 1.8, 3.7],
    )

    add_h1(doc, "3. Trạng thái loại xe trên app")
    add_table(
        doc,
        ["Trạng thái CMS", "Điều kiện thực tế", "App khách"],
        [
            ["Hoạt động", "Có bảng giá và có tài xế available", "Cho chọn"],
            ["Hoạt động", "Chưa có giá hoặc chưa có tài xế available", "Hiển thị Không khả dụng và disable"],
            ["Tạm dừng", "Không xét giá/tài xế", "Ẩn khỏi app"],
        ],
        [1.25, 3.0, 2.05],
    )
    add_body(doc, "Lưu ý: Hoạt động chỉ có nghĩa loại xe được phép kinh doanh. Khách có chọn được hay không còn phụ thuộc giá và tài xế available tại thời điểm booking.")

    add_h1(doc, "4. Quy trình CMS tạo loại xe và giá")
    add_step_flow(
        doc,
        "Flow tạo loại xe",
        [
            "Admin mở Dữ liệu vận tải → Tuyến & Lịch chạy → Loại xe.",
            "Admin thêm loại xe và chọn loại dịch vụ: Xe máy hoặc Xe hơi.",
            "Nhập mã, tên, phân loại xe, số ghế, trạng thái Hoạt động/Tạm dừng.",
            "Lưu loại xe. Nếu Hoạt động, loại xe sẵn sàng để tạo giá.",
        ],
    )
    add_step_flow(
        doc,
        "Flow tạo giá",
        [
            "Admin mở Tài chính → Giá tiền.",
            "Chọn tab Xe máy hoặc Xe hơi.",
            "Bấm Thêm bảng giá và chọn loại xe áp dụng.",
            "Nhập giá mở cửa, khung km, giá theo giờ/thời điểm nếu có.",
            "Lưu bảng giá. App khách chỉ cho chọn khi loại xe Hoạt động + có giá + có tài xế available.",
        ],
    )

    add_h1(doc, "5. Quy trình gán tài xế cho loại xe")
    add_table(
        doc,
        ["Trường", "Cách xử lý"],
        [
            ["Nhóm xe", "Xe máy hoặc Xe hơi"],
            ["Ô tô 6–7 chỗ", "Checkbox nhanh, chỉ áp dụng cho tài xế Xe hơi"],
            ["Premium", "Checkbox nhanh, áp dụng cho Bike/Car nếu tài xế đủ chuẩn"],
            ["Số ghế phục vụ khách", "Tự suy ra: Bike = 1, Car mặc định = 4, Car tick Ô tô 6–7 chỗ = 6"],
            ["Biển số", "Nhập theo xe riêng của tài xế"],
            ["Loại xe hệ thống cho phép nhận", "Hệ thống tự tính, admin chỉ xem"],
            ["Tùy chỉnh ngoại lệ", "Chỉ bật khi cần xử lý trường hợp đặc biệt"],
        ],
        [2.0, 4.3],
    )
    add_step_flow(
        doc,
        "Flow gán tài xế",
        [
            "Admin mở hồ sơ tài xế Bike/Car.",
            "Chọn Nhóm xe: Xe máy hoặc Xe hơi.",
            "Nếu là Xe hơi và tài xế chạy xe 6/7 chỗ, tick nhanh Ô tô 6–7 chỗ.",
            "Nếu tài xế đủ chuẩn premium, tick nhanh Premium.",
            "CMS tự suy ra số ghế phục vụ khách.",
            "CMS tự hiển thị danh sách loại xe hệ thống cho phép nhận.",
            "Nếu không có ngoại lệ, admin lưu hồ sơ.",
            "Nếu có ngoại lệ, admin bật Tùy chỉnh ngoại lệ rồi chọn lại danh sách loại xe được nhận.",
        ],
    )

    add_h1(doc, "6. Rule mặc định tài xế được nhận loại xe nào")
    add_table(
        doc,
        ["Nhóm xe/checkbox trong hồ sơ", "Premium", "Loại xe được nhận"],
        [
            ["Xe máy, không tick Premium", "Không", "Bike phổ thông"],
            ["Xe máy, tick Premium", "Có", "Bike Premium, Bike phổ thông"],
            ["Xe hơi, không tick Ô tô 6–7 chỗ", "Không", "Car 04 phổ thông"],
            ["Xe hơi, không tick Ô tô 6–7 chỗ", "Có", "Car 4 Premium, Car 04 phổ thông"],
            ["Xe hơi, tick Ô tô 6–7 chỗ", "Không", "Car 06 phổ thông, Car 04 phổ thông"],
            ["Xe hơi, tick Ô tô 6–7 chỗ", "Có", "Car 06 Premium, Car 06 phổ thông, Car 4 Premium, Car 04 phổ thông"],
        ],
        [2.15, 1.0, 3.15],
    )
    add_bullets(
        doc,
        [
            "Xe máy chỉ nhận loại xe máy; xe hơi chỉ nhận loại xe hơi.",
            "Mặc định tài xế Xe hơi là phổ thông 4 chỗ.",
            "Tài xế Xe hơi chạy xe 6/7 chỗ chỉ cần tick nhanh Ô tô 6–7 chỗ.",
            "BE xác định loại xe 6–7 chỗ bằng dữ liệu Loại xe: serviceType = CAR và seats >= 6 && seats <= 7, không hardcode theo tên loại xe.",
            "Loại Premium chỉ được nhận khi tick nhanh Premium.",
            "Nếu tạo loại xe mới nhưng chưa có rule riêng, tài xế mặc định chỉ nhận đúng loại xe đang chạy.",
            "Override chỉ dùng cho ngoại lệ, không dùng làm cách vận hành chính.",
        ],
    )

    add_h1(doc, "7. Hướng dẫn setup phân loại loại xe cho tài xế trên CMS")
    add_body(doc, "Admin không cần chọn từng loại xe con cho tài xế trong điều kiện bình thường. Chỉ cần setup theo checkbox nhanh.")
    add_body(doc, "Vị trí thao tác: Nhà xe & Tài xế → Tài xế Bike/Car → Thêm/Sửa tài xế")
    add_step_flow(
        doc,
        "Các bước setup",
        [
            "Chọn Nhóm xe: Xe máy hoặc Xe hơi.",
            "Nhập họ tên, số điện thoại, biển số và trạng thái ban đầu.",
            "Nếu là Xe hơi: không tick gì thêm thì hệ thống hiểu Car 04 phổ thông, số ghế = 4.",
            "Nếu là Xe hơi 6/7 chỗ: tick Ô tô 6–7 chỗ, hệ thống xét các loại CAR có seats trong khoảng 6–7.",
            "Nếu tài xế đủ chuẩn premium: tick Premium.",
            "Kiểm tra vùng Loại xe hệ thống cho phép nhận.",
            "Nếu danh sách đúng thì lưu hồ sơ.",
            "Chỉ bật Tùy chỉnh ngoại lệ khi BA/admin cần cấp khác rule mặc định.",
        ],
    )
    add_h2(doc, "Hình ảnh thao tác trên mock UI")
    add_image(doc, "driver-classification-bike-default.png", "Hình 4 — Tài xế Xe máy mặc định")
    add_body(doc, "Khi chọn Xe máy và không tick Premium, hệ thống tự hiểu tài xế nhận Bike phổ thông.")
    add_image(doc, "driver-classification-car-default.png", "Hình 5 — Tài xế Xe hơi mặc định 4 chỗ")
    add_body(doc, "Khi chọn Xe hơi và không tick Ô tô 6–7 chỗ hoặc Premium, hệ thống tự hiểu tài xế nhận Car 04 phổ thông, số ghế = 4.")
    add_image(doc, "driver-classification-car-6-7-premium.png", "Hình 6 — Tài xế Xe hơi 6–7 chỗ + Premium")
    add_body(doc, "Khi tick Ô tô 6–7 chỗ và Premium, hệ thống tự hiểu tài xế có năng lực nhận nhóm xe 6–7 chỗ và nhóm Premium. Danh sách loại xe được nhận vẫn dựa trên các Loại xe đang có trong hệ thống.")
    add_table(
        doc,
        ["Admin setup", "Hệ thống hiểu là", "Số ghế", "Loại xe được phép nhận"],
        [
            ["Xe máy, không tick Premium", "Bike phổ thông", "1", "Bike phổ thông"],
            ["Xe máy, tick Premium", "Bike Premium", "1", "Bike Premium; Bike phổ thông"],
            ["Xe hơi, không tick Ô tô 6–7 chỗ/Premium", "Car 04 phổ thông", "4", "Car 04 phổ thông"],
            ["Xe hơi, tick Premium", "Car 4 Premium", "4", "Car 4 Premium; Car 04 phổ thông"],
            ["Xe hơi, tick Ô tô 6–7 chỗ", "Car 06 phổ thông", "6", "Car 06 phổ thông; Car 04 phổ thông"],
            ["Xe hơi, tick Ô tô 6–7 chỗ + Premium", "Car 06 Premium", "6", "Car 06 Premium; Car 06 phổ thông; Car 4 Premium; Car 04 phổ thông"],
        ],
        [1.7, 1.35, 0.7, 2.55],
    )
    add_body(doc, "Nếu hệ thống có thêm Car 07 với seats = 7, tài xế tick Ô tô 6–7 chỗ cũng được xét nhận Car 07 theo cùng rule. Nếu chưa có loại CAR nào có seats từ 6 đến 7 đang hoạt động, CMS vẫn lưu năng lực tài xế nhưng preview chỉ hiển thị các loại CAR hiện có.")
    add_callout(
        doc,
        "Khi nào dùng Tùy chỉnh ngoại lệ?",
        "Không dùng override để setup 6 chỗ hoặc premium, vì hai trường hợp này đã có checkbox nhanh. Override chỉ dùng cho ngoại lệ như khóa tạm Premium, thử nghiệm loại xe đặc biệt, hoặc khóa một loại xe cụ thể theo chính sách vận hành.",
    )
    add_image(doc, "driver-classification-exception-override.png", "Hình 7 — Case ngoại lệ: admin chọn thủ công loại xe được nhận")
    add_body(doc, "Ví dụ: tài xế có xe 6–7 chỗ và đủ chuẩn Premium, nhưng trong thời gian thử việc admin chỉ cho nhận Car 04 phổ thông. Khi bật Tùy chỉnh ngoại lệ, matching dùng đúng danh sách được tick và CMS bắt buộc nhập Lý do ngoại lệ.")
    add_h2(doc, "Các trường hợp ngoại lệ thường gặp")
    add_table(
        doc,
        ["Trường hợp", "Cách setup trên CMS", "Kết quả mong muốn"],
        [
            ["Khóa Premium tạm thời", "Giữ tick Premium nếu hồ sơ vẫn đạt, bật Tùy chỉnh ngoại lệ và bỏ loại Premium", "Tài xế chỉ nhận nhóm phổ thông trong thời gian bị khóa"],
            ["Xe 6–7 chỗ nhưng chỉ chạy 4 chỗ khi thử việc", "Tick Ô tô 6–7 chỗ để lưu năng lực, bật Tùy chỉnh ngoại lệ và chỉ giữ nhóm Car 04 phù hợp", "Tài xế không nhận cuốc 6–7 chỗ"],
            ["Cấp thử loại xe mới chưa nằm trong rule", "Bật Tùy chỉnh ngoại lệ và chọn thêm loại xe thử nghiệm", "Matching đưa tài xế vào loại xe thử nghiệm"],
            ["Khóa một loại xe do giấy tờ/phương tiện không đạt", "Bật Tùy chỉnh ngoại lệ và bỏ loại xe đó", "Tài xế không nhận loại bị khóa, vẫn có thể nhận loại khác"],
            ["Chỉ nhận Premium hoặc chỉ nhận phổ thông", "Bật Tùy chỉnh ngoại lệ và chỉ tick đúng nhóm được phép", "Matching không tự cộng thêm nhóm còn lại"],
            ["Đổi xe tạm thời khác hồ sơ gốc", "Nếu ngắn hạn thì bật override và ghi rõ thời gian/lý do; nếu dài hạn thì cập nhật lại hồ sơ xe", "Tài xế chỉ nhận loại xe phù hợp xe đang chạy thực tế"],
            ["Loại xe master data nhập sai số ghế", "Không xử lý bằng override; sửa số ghế ở Loại xe", "BE phân nhóm đúng theo seats >= 6 && seats <= 7"],
            ["Chưa có loại CAR 6–7 chỗ đang hoạt động", "Không cần override; tick Ô tô 6–7 chỗ vẫn lưu năng lực", "Preview/matching chỉ dùng các loại xe hiện có"],
        ],
        [1.85, 2.6, 1.85],
    )
    add_bullets(
        doc,
        [
            "Override luôn ưu tiên hơn rule mặc định khi matching.",
            "Khi tắt override, hệ thống quay lại tính quyền nhận theo checkbox nhanh.",
            "Mỗi lần bật/tắt override hoặc thay danh sách override cần ghi audit log.",
            "CMS bắt buộc nhập Lý do ngoại lệ khi bật override.",
            "Nếu ngoại lệ có thời hạn, BA/admin nên ghi rõ mốc bắt đầu/kết thúc trong lý do.",
        ],
    )

    add_h1(doc, "8. Logic BE khi khách đặt Bike/Car")
    add_step_flow(
        doc,
        "Flow booking",
        [
            "FE lấy danh sách loại xe Bike/Car.",
            "Nếu loại xe Tạm dừng: app khách ẩn.",
            "Nếu loại xe Hoạt động nhưng chưa có giá hoặc không có tài xế available: app hiển thị Không khả dụng và disable.",
            "Nếu loại xe Hoạt động + có giá + có tài xế available: app cho khách chọn.",
            "Khách xác nhận booking và đi qua bước thanh toán/tạm giữ theo Epic 2.",
            "BE ghép tài xế theo đúng loại xe khách chọn.",
        ],
    )
    add_body(doc, "Khi ghép tài xế, BE chỉ lấy tài xế thỏa tất cả điều kiện:")
    add_bullets(
        doc,
        [
            "Cùng nhóm xe Bike hoặc Car với booking.",
            "Loại xe khách chọn nằm trong danh sách tài xế được phép nhận.",
            "Tài xế Online.",
            "Tài xế không Bận, không có chuyến đang chạy.",
            "Không có offer đang chờ accept.",
            "GPS/heartbeat còn hợp lệ.",
            "Hồ sơ và giấy tờ còn hiệu lực.",
            "Vị trí tài xế nằm trong bán kính tìm kiếm của loại xe.",
        ],
    )
    add_callout(
        doc,
        "Available không phải trạng thái thứ tư",
        "Online/Offline/Bận là trạng thái của tài xế. Available là kết quả tính tại thời điểm hệ thống tìm tài xế cho một booking cụ thể.",
    )

    add_h1(doc, "9. Logic xử lý override")
    add_body(doc, "Mặc định: allowedTypes = rule(nhóm xe, checkbox Ô tô 6–7 chỗ, checkbox Premium)")
    add_body(doc, "Khi bật override: allowedTypes = danh sách admin chọn riêng cho tài xế đó")
    add_bullets(
        doc,
        [
            "Chỉ bật override khi cần ngoại lệ nghiệp vụ.",
            "Khi tắt override, hệ thống quay lại rule mặc định.",
            "Audit log cần ghi lại ai bật/tắt override và danh sách loại xe sau khi lưu.",
            "BA cần xác nhận từng ngoại lệ để tránh cấp sai loại Premium.",
            "Nếu override đang bật, matching dùng đúng danh sách override admin đã chọn và không tự cộng thêm quyền từ checkbox Ô tô 6–7 chỗ/Premium.",
            "Nếu override đang bật mà thiếu Lý do ngoại lệ, CMS không cho lưu.",
        ],
    )

    add_h1(doc, "10. Checklist UAT nhanh")
    add_table(
        doc,
        ["Case", "Kỳ vọng"],
        [
            ["Xe máy, không tick Premium", "Chỉ nhận Bike phổ thông"],
            ["Xe máy, tick Premium", "Nhận Bike Premium và Bike phổ thông"],
            ["Xe hơi, không tick Ô tô 6–7 chỗ/Premium", "Chỉ nhận Car 04 phổ thông"],
            ["Xe hơi, tick Premium", "Nhận Car 4 Premium và Car 04 phổ thông"],
            ["Xe hơi, tick Ô tô 6–7 chỗ", "Nhận Car 06 phổ thông và Car 04 phổ thông"],
            ["Xe hơi, tick Ô tô 6–7 chỗ + Premium", "Nhận đủ nhóm Car 06/04 Premium/phổ thông theo rule"],
            ["Loại xe tạm dừng", "App khách ẩn loại xe"],
            ["Hoạt động nhưng chưa có tài xế available", "App khách hiển thị Không khả dụng và disable"],
            ["Driver Online nhưng sai loại xe", "Không được đưa vào matching"],
            ["Bật override cho một tài xế", "Matching dùng danh sách override thay vì rule mặc định"],
        ],
        [2.75, 3.55],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("HaHaGo Bike/Car CMS flow")
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
